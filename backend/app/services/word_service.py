import io
from datetime import datetime

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_student_profile_word(student_data: dict) -> bytes:
    buffer = io.BytesIO()
    
    if not HAS_DOCX:
        text_content = f"V.S.B. ENGINEERING COLLEGE (AUTONOMOUS)\nSTUDENT COMPLETE PROFILE\n\nName: {student_data.get('full_name')}\nReg No: {student_data.get('register_number')}\nDept: {student_data.get('department_name')}\nCGPA: {student_data.get('cgpa')}\n"
        buffer.write(text_content.encode('utf-8'))
        buffer.seek(0)
        return buffer.getvalue()

    doc = docx.Document()

    # Set Margins (0.6 in)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    # 1. Header Title
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_hdr.add_run("V.S.B. ENGINEERING COLLEGE (AUTONOMOUS)\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(178, 34, 34)  # Crimson red #B22222

    r_sub = p_hdr.add_run("Karur - 639 111, Tamil Nadu • Accredited by NAAC & NBA\n")
    r_sub.font.size = Pt(9.5)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)

    r_tag = p_hdr.add_run("OFFICIAL PERMANENT STUDENT DIGITAL PROFILE & ACADEMIC DOSSIER\n")
    r_tag.bold = True
    r_tag.font.size = Pt(11)
    r_tag.font.color.rgb = RGBColor(15, 23, 42)

    # Horizontal Divider Line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = RGBColor(178, 34, 34)
    p_div_run.bold = True

    # 2. General Information Table
    doc.add_heading("1. Student Demographic & Overview", level=2)
    
    table_info = doc.add_table(rows=7, cols=4)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.autofit = False

    info_fields = [
        ("Student Name:", student_data.get('full_name', '-'), "Register Number:", student_data.get('register_number', '-')),
        ("Roll Number:", student_data.get('roll_number', '-'), "Admission Number:", student_data.get('admission_number', '-')),
        ("Department:", student_data.get('department_name', '-'), "Batch / Year:", f"{student_data.get('batch', '-')} (Year {student_data.get('current_year', '-')})"),
        ("Current Semester:", f"Semester {student_data.get('current_semester', '-')}", "Section:", student_data.get('section_name', 'A')),
        ("Date of Birth / Gender:", f"{student_data.get('dob', '-')} ({student_data.get('gender', '-')})", "Blood Group:", student_data.get('blood_group', '-')),
        ("Contact Phone:", student_data.get('phone', '-'), "Email Address:", student_data.get('email', '-')),
        ("Father Name:", student_data.get('father_name', '-'), "Mother Name:", student_data.get('mother_name', '-'))
    ]

    for row_idx, data in enumerate(info_fields):
        row_cells = table_info.rows[row_idx].cells
        row_cells[0].paragraphs[0].add_run(data[0]).bold = True
        row_cells[1].paragraphs[0].add_run(str(data[1]))
        row_cells[2].paragraphs[0].add_run(data[2]).bold = True
        row_cells[3].paragraphs[0].add_run(str(data[3]))
        
        set_cell_background(row_cells[0], "F8FAFC")
        set_cell_background(row_cells[2], "F8FAFC")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. Academic Summary Table
    doc.add_heading("2. Academic Performance & CGPA Ledger", level=2)
    
    acad_table = doc.add_table(rows=2, cols=5)
    acad_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Cumulative CGPA", "Latest SGPA", "Attendance Rate", "Credits Earned", "Pending Arrears"]
    for idx, h in enumerate(headers):
        cell = acad_table.rows[0].cells[idx]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "B22222")

    values = [
        str(student_data.get('cgpa', '0.00')),
        str(student_data.get('sgpa', '0.00')),
        f"{student_data.get('attendance_percentage', '0')}%",
        str(student_data.get('credits_earned', '0')),
        str(student_data.get('arrears', '0'))
    ]

    for idx, v in enumerate(values):
        cell = acad_table.rows[1].cells[idx]
        p = cell.paragraphs[0]
        r = p.add_run(v)
        r.bold = True
        r.font.size = Pt(12)
        set_cell_background(cell, "F1F5F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. Certificates & Achievements
    doc.add_heading("3. Certificates & Achievements Ledger", level=2)
    certs = student_data.get('certificates', [])
    if certs and isinstance(certs, list):
        cert_table = doc.add_table(rows=len(certs) + 1, cols=5)
        cert_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_headers = ["#", "Certificate Name", "Type", "Organization / Issue Date", "Achievement / Number"]
        for idx, h in enumerate(c_headers):
            cell = cert_table.rows[0].cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            set_cell_background(cell, "E2E8F0")

        for r_idx, c in enumerate(certs):
            row_cells = cert_table.rows[r_idx + 1].cells
            row_cells[0].paragraphs[0].add_run(str(r_idx + 1))
            row_cells[1].paragraphs[0].add_run(c.get('name', 'Certificate'))
            row_cells[2].paragraphs[0].add_run(c.get('type', 'General'))
            row_cells[3].paragraphs[0].add_run(f"{c.get('issued_by', '-')} ({c.get('issue_date', '-')})")
            row_cells[4].paragraphs[0].add_run(f"{c.get('achievement', '-')} [{c.get('certificate_number', '-')}]")
    else:
        p_nocert = doc.add_paragraph()
        p_nocert.add_run("No custom achievement certificates uploaded.").italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Footer Details
    p_ftr = doc.add_paragraph()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ftr = p_ftr.add_run(f"\nGenerated Date: {datetime.now().strftime('%d %B %Y, %I:%M %p')}\nAuthorized Executive Seal — V.S.B. Engineering College")
    r_ftr.font.size = Pt(9)
    r_ftr.font.color.rgb = RGBColor(100, 116, 139)

    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
